from io import BytesIO
import os
import json
from docx import Document
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import openai
import requests
from dotenv import load_dotenv
import re
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import requests
from PIL import Image
import boto3
import time
from deep_translator import GoogleTranslator


load_dotenv()
my_key = os.getenv('OPENAI_API_KEY')

app = Flask(__name__)
CORS(app, origins=["https://smartchef-app.onrender.com"])
# CORS(app)

os.environ['OPENAI_API_KEY'] = my_key
client = openai.OpenAI()

my_model = "gpt-4o-mini"

@app.route("/")
def index():
    return "השרת פועל! 🍽️"

@app.route('/api/recipe/name', methods=['POST'])
def get_recipe_name():
    user_request = request.json.get('request', '').strip()

    # קרא שם מתכון מהמודל
    recipe_name = name_recipe_with_ai(user_request)
    if "Error" in recipe_name:
        return jsonify({"error": recipe_name}), 500

    return jsonify({"recipe_name": recipe_name})


@app.route('/api/recipe/file', methods=['POST'])
def get_recipe_file():
    if request.json is None:
        return jsonify({"error": "No JSON body provided"}), 400
    user_request = request.json.get('request', '').strip()
    recipe_name = request.json.get('recipe_name', '').strip()

    try:
        completion = client.chat.completions.create(
            model=my_model,
            temperature=1.2,
            messages=[
                {
                    "role": "system",
                    "content": f"""Give me a kosher recipe for the request.the language must be only in the user's request language. Include the recipe name: {recipe_name}, full ingredient list, and step-by-step instructions.The quantity of the ingredient will not be represented by numbers but by writing the quantity in words. Use line breaks only.The format must be:
                                Ingredients:
                                - ingredient 1
                                - ingredient 2
                                - ...
                                Instructions:
                                • Step 1
                                • Step 2
                                • ...
                                """
                },
                {
                    "role": "user",
                    "content": f"{user_request}"
                }
            ]
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    recipe_text = completion.choices[0].message.content

    # url = 'https://malismartchef.s3.us-east-1.amazonaws.com/recipes/template3.docx'
    url = 'https://malismartchef.s3.us-east-1.amazonaws.com/recipes/smartChef.docx'

    response = requests.get(url)
    if response.status_code == 200:
        with open('template.docx', 'wb') as f:
            f.write(response.content)

    doc = Document('template.docx')

    title = doc.add_paragraph(recipe_name)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.font.size = Pt(29)
    run.font.color.rgb = RGBColor(255, 183, 77)
    run.bold = True 

    lines = recipe_text.split("\n")
    ingredients = []
    instructions = []

    for line in lines:
        line = line.strip()
        if line.startswith("-"):
            ingredients.append(line[1:])
        # elif line and line.split('.')[0].isdigit():
        elif line.startswith("•") and line.endswith("."):
            instructions.append(line[1:-1])
        elif line.startswith("•"):
            instructions.append(line[1:])

    ingredients_text = "\n".join(ingredients)
    instructions_text = "\n".join(instructions)

    ingredients_name = doc.add_paragraph("רכיבים")
    ingredients_name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run2 = ingredients_name.runs[0]
    run2.font.size = Pt(26)
    run2.font.color.rgb = RGBColor(255, 149, 0)
    run2.bold = True

    ingredients_list = doc.add_paragraph(ingredients_text)
    ingredients_list.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run3 = ingredients_list.runs[0]
    run3.font.size = Pt(16)
    run3.font.color.rgb = RGBColor(255, 255, 255) 

    instructions_name = doc.add_paragraph("הוראות הכנה")
    instructions_name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run4 = instructions_name.runs[0]
    run4.font.size = Pt(26)
    run4.font.color.rgb = RGBColor(245, 124, 0)
    run4.bold = True

    instructions_list = doc.add_paragraph(instructions_text)
    instructions_list.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run5 = instructions_list.runs[0]
    run5.font.size = Pt(16)
    run5.font.color.rgb = RGBColor(255, 255, 255) 

    # doc.add_paragraph(recipe_name)
    # doc.add_paragraph("רכיבים:")
    # doc.add_paragraph(ingredients_text)
    # doc.add_paragraph("הוראות הכנה:")
    # doc.add_paragraph(instructions_text)

    # for para in doc.paragraphs:
    #     if "{{recipe_name}}" in para.text:
    #         para.text = para.text.replace("{{recipe_name}}", recipe_name)
    #     if "{{ingredients}}" in para.text:
    #         para.text = para.text.replace("{{ingredients}}", ingredients_text)
    #     if "{{instructions}}" in para.text:
    #         para.text = para.text.replace("{{instructions}}", instructions_text)

    # for ingredient_line in ingredients:
    #     # הסרת מקף ורווחים מיותרים מההתחלה
    #     cleaned_line = ingredient_line.lstrip('- ').strip()
    #     match = re.match(r'^(\d+(?:/\d+)?)\s*(.*)', cleaned_line)
    #     if match:
    #         quantity = match.group(1).strip()
    #         item = match.group(2).strip()
    #         print(f"{quantity} {item}\n {item} {quantity}")
    #         ingredient_text = f"- {quantity} {item}"
    #     else:
    #         # אם לא נמצאה כמות בהתחלה אחרי הסרת המקף, נוסיף מקף והשורה המקורית
    #         ingredient_text = "- " + cleaned_line
    #     doc.add_paragraph(ingredient_text)

    # # הוספת הוראות
    # for instruction_line in instructions:
    #     match = re.match(r'^(.*?)(\d+)\.(.*)$', instruction_line)
    #     if match:
    #         instruction_text = f"{match.group(2)}. {match.group(1).strip()} {match.group(3).strip()}"
    #     else:
    #         instruction_text = instruction_line.strip()
    #     doc.add_paragraph(instruction_text)
    

    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    output_filename = f"{recipe_name}.docx"
    print(output_stream)
    return send_file(output_stream, as_attachment=True, download_name=output_filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


def name_recipe_with_ai(user_request):
    try:
        completion = client.chat.completions.create(
            model=my_model,
            messages=[
                {
                    "role": "system",
                    # "content": f"""You are a precise chef who names recipes. The recipe name must be in the language of the user's request. Given the following request, return a recipe name that matches the request, a generic name without events, dates, and specific descriptions from the request."""
                    "content": f"""You are a precise chef who names recipes. The recipe name must be in the language of the user's request. Given the following request, return a recipe name that matches the request, a generic name without events, dates, and specific descriptions from the request. Address allergies and preferences only if it does not conflict with the user's request."""
                },
                {
                    "role": "user",
                    "content": f"""{user_request}"""
                }
            ]
        )

        answer = completion.choices[0].message.content
        return answer

    except Exception as e:
        return f"Error: {e}"
    
# @app.route('/api/recipe/image', methods=['POST'])
# def generate_image_with_dalle():
#     data = request.get_json()
#     prompt = data.get('prompt') 

#     if not prompt:
#         return {"error": "Prompt is required"}, 400 

#     output_path = "dalle_output.png"
#     headers = {
#         "Authorization": f"Bearer {my_key}",
#         "Content-Type": "application/json",
#     }
#     data = {
#         "prompt": prompt,
#         "n": 1,
#         "size": "512x512"
#     }
#     response = requests.post("https://api.openai.com/v1/images/generations", headers=headers, json=data)
#     response.raise_for_status()
#     image_url = response.json()["data"][0]["url"]
#     print(image_url)
#     image_data = requests.get(image_url).content
#     image = Image.open(BytesIO(image_data))
#     image.save(output_path)
#     return output_path


@app.route('/api/recipe/image', methods=['POST'])
def generate_image_with_dalle():
    data = request.get_json()
    prompt = data.get('prompt')

    if not prompt:
        return {"error": "Prompt is required"}, 400
    
    prompt_en = GoogleTranslator(source='auto', target='en').translate(prompt)

    headers = {
        "Authorization": f"Bearer {my_key}",
        "Content-Type": "application/json",
    }
    data = {
        "prompt": prompt_en,
        "n": 1,
        "size": "512x512"
    }
    response = requests.post("https://api.openai.com/v1/images/generations", headers=headers, json=data)
    response.raise_for_status()
    image_url = response.json()["data"][0]["url"]
    print(image_url)
    image_data = requests.get(image_url).content
    image = Image.open(BytesIO(image_data))

    img_byte_arr = BytesIO()
    image.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    img_byte_arr.seek(0)
    with open("temp_image.png", "wb") as f:
        f.write(img_byte_arr.getvalue())
    print(type(img_byte_arr))

    return send_file(img_byte_arr, mimetype='image/png')


if __name__ == '__main__':
    app.run(debug=True)

